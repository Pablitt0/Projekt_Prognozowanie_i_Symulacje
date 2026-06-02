# -*- coding: utf-8 -*-
"""Konwertuje raport.md -> raport.docx"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(SCRIPT_DIR, "raport.md")
DST = os.path.join(SCRIPT_DIR, "raport.docx")

BLUE = RGBColor(0x1a, 0x5c, 0x96)

doc = Document()

# ── Marginesy ────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Style normalne ───────────────────────────────────────────
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)

def set_heading_style(para, level):
    sizes = {1: 18, 2: 15, 3: 13}
    run = para.runs[0] if para.runs else para.add_run(para.text)
    run.font.bold = True
    run.font.color.rgb = BLUE
    run.font.size = Pt(sizes.get(level, 12))

def apply_inline(para, text):
    """Obsługuje **bold**, `code` i zwykły tekst."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(2):  # bold
            r = para.add_run(m.group(2))
            r.bold = True
        elif m.group(3):  # code
            r = para.add_run(m.group(3))
            r.font.name = "Courier New"
            r.font.size = Pt(10)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])

def add_table(doc, rows):
    """rows[0] = nagłówki, rows[1:] = dane."""
    if not rows:
        return
    cols = len(rows[0])
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = "Table Grid"
    for ci, hdr in enumerate(rows[0]):
        cell = tbl.rows[0].cells[ci]
        cell.text = ""
        run = cell.paragraphs[0].add_run(hdr.strip())
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # tło nagłówka
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1A5C96")
        tcPr.append(shd)
    for ri, row in enumerate(rows[1:], start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            apply_inline(cell.paragraphs[0], val.strip())

# ── Parser ───────────────────────────────────────────────────
with open(SRC, encoding="utf-8") as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip("\n")

    # Nagłówki
    if line.startswith("### "):
        p = doc.add_heading(line[4:], level=3)
        set_heading_style(p, 3)
        i += 1; continue
    if line.startswith("## "):
        p = doc.add_heading(line[3:], level=2)
        set_heading_style(p, 2)
        i += 1; continue
    if line.startswith("# "):
        p = doc.add_heading(line[2:], level=1)
        set_heading_style(p, 1)
        i += 1; continue

    # Separator ---
    if re.match(r"^-{3,}$", line):
        doc.add_paragraph("─" * 60)
        i += 1; continue

    # Tabela
    if line.startswith("|"):
        tbl_lines = []
        while i < len(lines) and lines[i].startswith("|"):
            tbl_lines.append(lines[i].rstrip("\n"))
            i += 1
        rows = []
        for tl in tbl_lines:
            cells = [c.strip() for c in tl.strip("|").split("|")]
            if all(re.match(r"^[-: ]+$", c) for c in cells):
                continue
            rows.append(cells)
        add_table(doc, rows)
        doc.add_paragraph()
        continue

    # Blok kodu
    if line.startswith("```"):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].startswith("```"):
            code_lines.append(lines[i].rstrip("\n"))
            i += 1
        i += 1
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Cm(1)
        continue

    # Cytat >
    if line.startswith("> "):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(1)
        apply_inline(p, line[2:])
        run = p.runs[0] if p.runs else None
        if run:
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            run.italic = True
        i += 1; continue

    # Lista punktowana
    if re.match(r"^(\d+\.|[-*]) ", line):
        p = doc.add_paragraph(style="List Bullet")
        content = re.sub(r"^(\d+\.|[-*]) ", "", line)
        apply_inline(p, content)
        i += 1; continue

    # Pusta linia
    if line.strip() == "":
        i += 1; continue

    # Zwykły akapit
    p = doc.add_paragraph(style="Normal")
    apply_inline(p, line)
    i += 1

doc.save(DST)
print(f"Zapisano: {DST}")
